"""Template engine service — Section 7 of the technical specification.

Handles template CRUD, auto-population of fields from project/researcher
data, conditional section evaluation, and DOCX document generation using
python-docx.
"""

import io
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.enums import (
    GeneratedDocumentStatus,
    PersonnelTemplateType,
    TemplateCategory,
)
from app.models.project import Project
from app.models.researcher import Researcher
from app.models.template import DocumentTemplate, GeneratedDocument
from app.schemas.template import (
    TemplateFieldPreview,
    TemplatePreviewResponse,
)
from app.templates.ec_consortium_definitions import (
    EC_CONSORTIUM_TEMPLATE_DEFINITIONS,
)
from app.templates.mission_travel_definitions import (
    MISSION_TRAVEL_TEMPLATE_DEFINITIONS,
)
from app.templates.personnel_definitions import (
    PERSONNEL_TEMPLATE_DEFINITIONS,
)
from app.templates.procurement_definitions import (
    PROCUREMENT_TEMPLATE_DEFINITIONS,
)
from app.templates.reporting_definitions import (
    REPORTING_TEMPLATE_DEFINITIONS,
)

# Base directory for generated documents
GENERATED_DOCS_DIR = Path("generated_documents")


# --- Template CRUD ---


async def seed_personnel_templates(db: AsyncSession) -> list[DocumentTemplate]:
    """Seed all personnel template definitions into the database.

    Idempotent: skips templates that already exist (by slug).
    """
    return await _seed_definitions(db, PERSONNEL_TEMPLATE_DEFINITIONS)


async def seed_all_templates(db: AsyncSession) -> list[DocumentTemplate]:
    """Seed all template definitions (all categories) into the database.

    Idempotent: skips templates that already exist (by slug).
    """
    all_defs: dict[str, dict] = {}
    # Personnel templates are keyed by enum, convert to slug-keyed
    for _ptype, defn in PERSONNEL_TEMPLATE_DEFINITIONS.items():
        all_defs[defn["slug"]] = defn
    # Other categories are already slug-keyed
    all_defs.update(EC_CONSORTIUM_TEMPLATE_DEFINITIONS)
    all_defs.update(REPORTING_TEMPLATE_DEFINITIONS)
    all_defs.update(MISSION_TRAVEL_TEMPLATE_DEFINITIONS)
    all_defs.update(PROCUREMENT_TEMPLATE_DEFINITIONS)
    return await _seed_definitions(db, all_defs)


async def _seed_definitions(
    db: AsyncSession, definitions: dict
) -> list[DocumentTemplate]:
    """Seed template definitions into the database.

    Idempotent: skips templates that already exist (by slug).
    """
    created: list[DocumentTemplate] = []
    for _key, defn in definitions.items():
        existing = await get_template_by_slug(db, defn["slug"])
        if existing:
            continue
        template = DocumentTemplate(
            name=defn["name"],
            slug=defn["slug"],
            description=defn["description"],
            category=defn["category"],
            personnel_type=defn.get("personnel_type"),
            version=1,
            field_definitions=defn["field_definitions"],
            conditional_sections=defn.get("conditional_sections"),
        )
        db.add(template)
        created.append(template)
    if created:
        await db.flush()
        for t in created:
            await db.refresh(t)
    return created


async def get_template_by_slug(
    db: AsyncSession, slug: str
) -> DocumentTemplate | None:
    """Get a template by its slug."""
    stmt = select(DocumentTemplate).where(DocumentTemplate.slug == slug)
    return (await db.execute(stmt)).scalar_one_or_none()


async def get_template(
    db: AsyncSession, template_id: uuid.UUID
) -> DocumentTemplate:
    """Get a template by ID or raise 404."""
    stmt = select(DocumentTemplate).where(
        DocumentTemplate.id == template_id
    )
    result = (await db.execute(stmt)).scalar_one_or_none()
    if not result:
        raise HTTPException(status_code=404, detail="Template not found")
    return result


async def list_templates(
    db: AsyncSession,
    category: TemplateCategory | None = None,
) -> list[DocumentTemplate]:
    """List all templates, optionally filtered by category."""
    stmt = select(DocumentTemplate).order_by(
        DocumentTemplate.category, DocumentTemplate.name
    )
    if category:
        stmt = stmt.where(DocumentTemplate.category == category)
    return list((await db.execute(stmt)).scalars().all())


# --- Field Auto-Population ---


async def _get_project_data(
    db: AsyncSession, project_id: uuid.UUID
) -> dict[str, str | None]:
    """Extract auto-population data from a project."""
    stmt = select(Project).where(
        Project.id == project_id, Project.deleted_at.is_(None)
    )
    project = (await db.execute(stmt)).scalar_one_or_none()
    if not project:
        return {}
    return {
        "project.acronym": project.acronym,
        "project.full_title": project.full_title,
        "project.grant_agreement_number": project.grant_agreement_number,
        "project.programme": project.programme.value if project.programme else None,
        "project.cost_model": project.cost_model.value if project.cost_model else None,
        "project.role": project.role.value if project.role else None,
        "project.start_date": (
            project.start_date.isoformat() if project.start_date else None
        ),
        "project.end_date": (
            project.end_date.isoformat() if project.end_date else None
        ),
        "project.duration_months": (
            str(project.duration_months)
            if project.duration_months
            else None
        ),
        "project.total_budget": (
            str(project.total_budget) if project.total_budget else None
        ),
        "project.eu_contribution": (
            str(project.eu_contribution) if project.eu_contribution else None
        ),
        "project.funding_rate": (
            str(project.funding_rate) if project.funding_rate else None
        ),
        "project.internal_cost_center": project.internal_cost_center,
        "project.ec_project_officer": project.ec_project_officer,
        "project.status": project.status.value if project.status else None,
    }


async def _get_researcher_data(
    db: AsyncSession, researcher_id: uuid.UUID
) -> dict[str, str | None]:
    """Extract auto-population data from a researcher."""
    stmt = select(Researcher).where(
        Researcher.id == researcher_id,
        Researcher.deleted_at.is_(None),
    )
    researcher = (await db.execute(stmt)).scalar_one_or_none()
    if not researcher:
        return {}
    return {
        "researcher.name": researcher.name,
        "researcher.email": researcher.email,
        "researcher.position": (
            researcher.position.value if researcher.position else None
        ),
        "researcher.contract_type": (
            researcher.contract_type.value
            if researcher.contract_type
            else None
        ),
        "researcher.fte": str(researcher.fte) if researcher.fte else None,
        "researcher.annual_gross_cost": (
            str(researcher.annual_gross_cost)
            if researcher.annual_gross_cost
            else None
        ),
        "researcher.hourly_rate": (
            str(researcher.hourly_rate) if researcher.hourly_rate else None
        ),
        "researcher.productive_hours": str(researcher.productive_hours),
        "researcher.start_date": (
            researcher.start_date.isoformat()
            if researcher.start_date
            else None
        ),
        "researcher.end_date": (
            researcher.end_date.isoformat()
            if researcher.end_date
            else None
        ),
    }


def _resolve_auto_fields(
    field_defs: list[dict],
    project_data: dict[str, str | None],
    researcher_data: dict[str, str | None],
) -> dict[str, str | None]:
    """Resolve auto-populated field values from data sources."""
    values: dict[str, str | None] = {}
    all_data = {**project_data, **researcher_data}

    for field in field_defs:
        if field.get("field_type") == "auto":
            source = field.get("data_source", "")
            values[field["name"]] = all_data.get(source)
    return values


def _evaluate_conditional_sections(
    conditions: dict | None,
    project_data: dict[str, str | None],
) -> list[str]:
    """Return list of active conditional section names."""
    if not conditions:
        return []

    active: list[str] = []
    for key, cond in conditions.items():
        field = cond.get("condition_field", "")
        expected = cond.get("condition_value", "")
        operator = cond.get("condition_operator", "eq")
        actual = project_data.get(field)

        if operator == "eq" and actual == expected:
            active.append(key)
        elif operator == "neq" and actual != expected:
            active.append(key)
        elif operator == "in" and actual in expected.split(","):
            active.append(key)
    return active


async def preview_template(
    db: AsyncSession,
    template_id: uuid.UUID,
    project_id: uuid.UUID | None = None,
    researcher_id: uuid.UUID | None = None,
) -> TemplatePreviewResponse:
    """Preview a template with auto-populated fields filled in."""
    template = await get_template(db, template_id)

    project_data: dict[str, str | None] = {}
    researcher_data: dict[str, str | None] = {}
    if project_id:
        project_data = await _get_project_data(db, project_id)
    if researcher_id:
        researcher_data = await _get_researcher_data(db, researcher_id)

    field_defs = template.field_definitions.get("fields", [])
    auto_values = _resolve_auto_fields(
        field_defs, project_data, researcher_data
    )

    fields: list[TemplateFieldPreview] = []
    for fdef in field_defs:
        value = auto_values.get(fdef["name"])
        fields.append(
            TemplateFieldPreview(
                name=fdef["name"],
                label=fdef.get("label", fdef["name"]),
                field_type=fdef.get("field_type", "manual"),
                data_type=fdef.get("data_type", "text"),
                value=value,
                required=fdef.get("required", True),
                help_text=fdef.get("help_text"),
            )
        )

    active_sections = _evaluate_conditional_sections(
        template.conditional_sections, project_data
    )

    return TemplatePreviewResponse(
        template_id=template.id,
        template_name=template.name,
        template_version=template.version,
        fields=fields,
        conditional_sections_active=active_sections,
    )


# --- DOCX Generation ---


def _generate_docx(
    template: DocumentTemplate,
    all_values: dict[str, str | None],
    active_sections: list[str],
) -> io.BytesIO:
    """Generate a DOCX document from template and field values.

    Uses python-docx to create a structured document with:
    - Title and metadata
    - Auto-populated and manual field values in a structured layout
    - Active conditional sections included
    - Signature blocks
    """
    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading(template.name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata paragraph
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {template.version}").italic = True
    meta.add_run(
        f"  |  Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).italic = True

    doc.add_paragraph()  # spacer

    # --- Project Information Section ---
    doc.add_heading("Project Information", level=2)
    _add_field_table(
        doc,
        all_values,
        [
            ("Project Acronym", "project_acronym"),
            ("Full Title", "project_title"),
            ("Grant Agreement No.", "grant_agreement_number"),
            ("Programme", "programme"),
            ("Start Date", "project_start_date"),
            ("End Date", "project_end_date"),
            ("Cost Center", "cost_center"),
        ],
    )

    # --- Researcher Information Section ---
    doc.add_heading("Researcher Information", level=2)
    _add_field_table(
        doc,
        all_values,
        [
            ("Name", "researcher_name"),
            ("Email", "researcher_email"),
            ("Position", "position"),
            ("FTE", "fte"),
            ("Contract Start", "contract_start"),
            ("Contract End", "contract_end"),
            ("Annual Gross Cost", "annual_gross_cost"),
        ],
    )

    # --- Template-Specific Fields ---
    field_defs = template.field_definitions.get("fields", [])
    manual_fields = [
        f for f in field_defs if f.get("field_type") == "manual"
    ]
    if manual_fields:
        doc.add_heading("Contract Details", level=2)
        rows = [
            (f.get("label", f["name"]), f["name"])
            for f in manual_fields
        ]
        _add_field_table(doc, all_values, rows)

    # --- Conditional Sections ---
    if template.conditional_sections:
        for key, cond in template.conditional_sections.items():
            if key in active_sections:
                doc.add_heading(cond.get("section_name", key), level=2)
                doc.add_paragraph(cond.get("content_label", ""))

    # --- Signature Block ---
    doc.add_paragraph()  # spacer
    doc.add_heading("Signatures", level=2)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    for i, (role, line) in enumerate(
        [
            ("Principal Investigator", "Date: _______________"),
            ("Researcher", "Date: _______________"),
            ("Institutional Representative", "Date: _______________"),
        ]
    ):
        sig_table.cell(i, 0).text = f"{role}\n\n\nSignature: _______________"
        sig_table.cell(i, 1).text = line

    # Set document properties
    props = doc.core_properties
    props.author = "EU Project Management System"
    props.title = template.name
    props.subject = f"Template: {template.slug}, Version: {template.version}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_field_table(
    doc: DocxDocument,
    values: dict[str, str | None],
    rows: list[tuple[str, str]],
) -> None:
    """Add a two-column field-value table to the document."""
    filtered = [
        (label, key) for label, key in rows
        if values.get(key) is not None
    ]
    if not filtered:
        doc.add_paragraph("(No data available)")
        return

    table = doc.add_table(rows=len(filtered), cols=2)
    table.style = "Light List Accent 1"

    for i, (label, key) in enumerate(filtered):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(values.get(key, ""))


# --- Document Generation Workflow ---


async def generate_document(
    db: AsyncSession,
    template_id: uuid.UUID,
    project_id: uuid.UUID | None = None,
    researcher_id: uuid.UUID | None = None,
    manual_fields: dict | None = None,
    generated_by: str | None = None,
) -> GeneratedDocument:
    """Generate a DOCX document from a template.

    1. Loads template definition
    2. Auto-populates fields from project/researcher data
    3. Merges manual input fields
    4. Evaluates conditional sections
    5. Generates DOCX via python-docx
    6. Saves file and creates database record
    """
    template = await get_template(db, template_id)

    # Gather auto-populated data
    project_data: dict[str, str | None] = {}
    researcher_data: dict[str, str | None] = {}
    if project_id:
        project_data = await _get_project_data(db, project_id)
    if researcher_id:
        researcher_data = await _get_researcher_data(db, researcher_id)

    # Resolve auto fields
    field_defs = template.field_definitions.get("fields", [])
    auto_values = _resolve_auto_fields(
        field_defs, project_data, researcher_data
    )

    # Merge manual fields
    all_values = {**auto_values, **(manual_fields or {})}

    # Evaluate conditional sections
    active_sections = _evaluate_conditional_sections(
        template.conditional_sections, project_data
    )

    # Generate DOCX
    docx_buf = _generate_docx(template, all_values, active_sections)

    # Determine file name
    parts = [template.slug]
    if project_id and project_data.get("project.acronym"):
        parts.append(project_data["project.acronym"])
    if researcher_id and researcher_data.get("researcher.name"):
        name_slug = (
            researcher_data["researcher.name"]
            .replace(" ", "_")
            .lower()[:30]
        )
        parts.append(name_slug)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    parts.append(timestamp)
    file_name = "_".join(parts) + ".docx"

    # Save to filesystem
    GENERATED_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    file_path = GENERATED_DOCS_DIR / file_name
    file_path.write_bytes(docx_buf.getvalue())

    # Create database record
    now = datetime.now(timezone.utc)
    doc_record = GeneratedDocument(
        template_id=template.id,
        project_id=project_id,
        researcher_id=researcher_id,
        template_version=template.version,
        generated_by=generated_by,
        generated_at=now,
        input_data=all_values,
        file_path=str(file_path),
        file_name=file_name,
        status=GeneratedDocumentStatus.DRAFT,
    )
    db.add(doc_record)
    await db.flush()
    await db.refresh(doc_record)
    return doc_record


async def generate_document_bytes(
    db: AsyncSession,
    template_id: uuid.UUID,
    project_id: uuid.UUID | None = None,
    researcher_id: uuid.UUID | None = None,
    manual_fields: dict | None = None,
) -> tuple[io.BytesIO, str]:
    """Generate DOCX bytes without saving to filesystem.

    Returns (bytes_buffer, filename) for streaming download.
    """
    template = await get_template(db, template_id)

    project_data: dict[str, str | None] = {}
    researcher_data: dict[str, str | None] = {}
    if project_id:
        project_data = await _get_project_data(db, project_id)
    if researcher_id:
        researcher_data = await _get_researcher_data(db, researcher_id)

    field_defs = template.field_definitions.get("fields", [])
    auto_values = _resolve_auto_fields(
        field_defs, project_data, researcher_data
    )
    all_values = {**auto_values, **(manual_fields or {})}
    active_sections = _evaluate_conditional_sections(
        template.conditional_sections, project_data
    )

    docx_buf = _generate_docx(template, all_values, active_sections)

    parts = [template.slug]
    if project_data.get("project.acronym"):
        parts.append(project_data["project.acronym"])
    file_name = "_".join(parts) + ".docx"

    return docx_buf, file_name


# --- Generated Document CRUD ---


async def get_generated_document(
    db: AsyncSession, doc_id: uuid.UUID
) -> GeneratedDocument:
    """Get a generated document by ID."""
    stmt = select(GeneratedDocument).where(
        GeneratedDocument.id == doc_id
    )
    result = (await db.execute(stmt)).scalar_one_or_none()
    if not result:
        raise HTTPException(
            status_code=404, detail="Generated document not found"
        )
    return result


async def list_generated_documents(
    db: AsyncSession,
    template_id: uuid.UUID | None = None,
    project_id: uuid.UUID | None = None,
) -> list[GeneratedDocument]:
    """List generated documents with optional filters."""
    stmt = select(GeneratedDocument).order_by(
        GeneratedDocument.generated_at.desc()
    )
    if template_id:
        stmt = stmt.where(GeneratedDocument.template_id == template_id)
    if project_id:
        stmt = stmt.where(GeneratedDocument.project_id == project_id)
    return list((await db.execute(stmt)).scalars().all())


async def update_document_status(
    db: AsyncSession,
    doc_id: uuid.UUID,
    status: GeneratedDocumentStatus,
) -> GeneratedDocument:
    """Update the status of a generated document."""
    doc = await get_generated_document(db, doc_id)
    doc.status = status
    await db.flush()
    await db.refresh(doc)
    return doc
